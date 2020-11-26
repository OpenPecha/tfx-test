from docx import Document  # package name: python-docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
import re


def create_total_docx(chunks, path):
    # CONFIG
    indentation = 0.5
    tibetan_style = {
        'font': 'Jomolhari',
        'size': 11
    }
    pedurma_style = {
        'color': (112, 128, 144),
        'font': 'Jomolhari',
        'size': 8
    }
    semantic_style = {
        'font': 'Free Mono',
        'size': 10
    }
    communicative_style = {
        'font': 'Gentium',
        'size': 12
    }

    document = Document()
    styles = document.styles

    # TIBETAN
    bo_style = styles.add_style('Tibetan', WD_STYLE_TYPE.CHARACTER)
    bo_font = bo_style.font
    bo_font.name = tibetan_style['font']
    bo_font.size = Pt(tibetan_style['size'])
    # PEYDURMA NOTES
    note_style = styles.add_style('Peydurma Notes', WD_STYLE_TYPE.CHARACTER)
    note_font = note_style.font
    note_font.name = pedurma_style['font']
    # note_font.size = Pt(pedurma_style['size'])
    note_font.subscript = True
    c = pedurma_style['color']
    note_font.color.rgb = RGBColor(c[0], c[1], c[2])

    # COMMUNICATIVE VERSION
    com_style = styles.add_style('Communicative', WD_STYLE_TYPE.CHARACTER)
    com_font = com_style.font
    com_font.name = communicative_style['font']
    com_font.size = Pt(communicative_style['size'])

    # SEMANTIC VERSION
    sem_style = styles.add_style('Semantic', WD_STYLE_TYPE.CHARACTER)
    sem_style.base_style = styles['Normal']
    sem_font = sem_style.font
    sem_font.name = semantic_style['font']
    sem_font.size = Pt(semantic_style['size'])

    # COMMUNICATIVE PARAGRAPH
    com_par_style = styles.add_style('Com. paragraph', WD_STYLE_TYPE.PARAGRAPH)
    com_par_style.paragraph_format.space_before = Cm(0)
    com_par_style.paragraph_format.space_after = Cm(0)

    # OTHER PARAGRAPH
    other_par_style = styles.add_style('Other paragraph', WD_STYLE_TYPE.PARAGRAPH)
    other_par_style.paragraph_format.space_before = Cm(0)
    other_par_style.paragraph_format.space_after = Cm(1)
    other_par_style.paragraph_format.left_indent = Cm(indentation)
    other_par_style.paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE

    for chunk in chunks:
        com, others = chunk

        #########################
        # Communicative
        com_p = document.add_paragraph()
        com_p.style = 'Com. paragraph'
        com = re.split('(/[^/]+/)', com)
        for c in com:
            com_run = com_p.add_run('')
            com_run.style = 'Communicative'

            if c.startswith('/'):
                com_run.text = c[1:-1]
                com_run.italic = True
            else:
                com_run.text = c

        # ************************

        p = document.add_paragraph()
        p.style = 'Other paragraph'

        for pair in others:
            bo, sem = pair
            bo = re.split('(<.+?>)', bo)
            for b in bo:
                run = p.add_run(b)
                if b.startswith('<'):
                    run.style = 'Peydurma Notes'
                else:
                    run.style = 'Tibetan'
            p.add_run().add_break()

            run = p.add_run(sem)
            run.style = 'Semantic'
            run.add_break()

        # removing trailing newline
        p.runs[-1].text = p.runs[-1].text.rstrip('\n')
        # print('ok')
    out_path = path.parent / (path.stem + '.docx')
    document.save(str(out_path))


def create_trans_docx(pars, path):
    # CONFIG
    communicative_style = {
        'font': 'Gentium',
        'size': 12
    }

    document = Document()
    styles = document.styles

    # COMMUNICATIVE VERSION
    com_style = styles.add_style('Communicative', WD_STYLE_TYPE.CHARACTER)
    com_font = com_style.font
    com_font.name = communicative_style['font']
    com_font.size = Pt(communicative_style['size'])

    # COMMUNICATIVE PARAGRAPH
    com_par_style = styles.add_style('Com. paragraph trans', WD_STYLE_TYPE.PARAGRAPH)
    com_format = com_par_style.paragraph_format
    com_format.first_line_indent = Cm(1)
    # com_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_HI

    for par in pars.split('\n\n'):

        #########################
        # Communicative
        com_p = document.add_paragraph()
        com_p.style = 'Com. paragraph trans'

        com = re.split('(/[^/]+/)', par)
        for c in com:
            com_run = com_p.add_run('')
            com_run.style = 'Communicative'

            if c.startswith('/'):
                com_run.text = c[1:-1]
                com_run.italic = True
            else:
                com_run.text = c

    out_path = path.parent / (path.stem + '.docx')
    document.save(str(out_path))
